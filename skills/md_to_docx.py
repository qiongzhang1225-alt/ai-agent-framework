def md_to_docx(markdown_text: str, output_path: str = "") -> str:
    """把 Markdown 文本转成格式化的 Word (.docx) 文档。

    支持：
    - # ~ ###### 标题
    - **加粗** *斜体* `行内代码`
    - 有序 / 无序列表
    - | 表格 |
    - ```代码块```
    - ![图片](url)
    - $$ 数学公式 $$（转图片嵌入）
    - > 引用块

    参数：
        markdown_text: 完整 Markdown 源码
        output_path: 输出的 .docx 路径。空时自动生成到工作目录

    返回：生成的 .docx 文件路径
    """
    import os, re, uuid
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 样式设置 ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 7):
        hs = doc.styles[f"Heading {level}"]
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)

    # ── 行内样式解析 ──
    def _add_run(paragraph, text):
        """解析加粗/斜体/行内代码并添加到段落。"""
        # 按 ** ** * * ` ` 分段
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for p in parts:
            if p.startswith("**") and p.endswith("**"):
                run = paragraph.add_run(p[2:-2])
                run.bold = True
            elif p.startswith("*") and p.endswith("*") and not p.startswith("**"):
                run = paragraph.add_run(p[1:-1])
                run.italic = True
            elif p.startswith("`") and p.endswith("`"):
                run = paragraph.add_run(p[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            else:
                paragraph.add_run(p)

    # ── 公式转图片（用 matplotlib 渲染 LaTeX 公式）──
    def _render_formula(formula, output_dir, idx):
        """将 LaTeX 公式渲染为 PNG 图片。"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            fig, ax = plt.subplots(figsize=(len(formula) * 0.08 + 1, 0.6))
            ax.text(0.5, 0.5, f"${formula}$", fontsize=14,
                    ha="center", va="center", usetex=False)
            ax.axis("off")
            path = os.path.join(output_dir, f"formula_{idx}.png")
            plt.savefig(path, dpi=150, bbox_inches="tight", pad_inches=0.05)
            plt.close(fig)
            return path
        except Exception:
            return None

    md = markdown_text
    output_dir = os.path.dirname(output_path) if output_path else os.getcwd()
    os.makedirs(output_dir, exist_ok=True)
    formula_idx = [0]
    i = 0
    lines = md.split("\n")

    while i < len(lines):
        line = lines[i]

        # ── 跳过空行 ──
        if not line.strip():
            i += 1
            continue

        # ── 代码块 ──
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # ── 数学公式块 ──
        if line.strip().startswith("$$"):
            formula_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                formula_lines.append(lines[i])
                i += 1
            i += 1  # skip closing $$
            formula_text = " ".join(formula_lines)
            if formula_text:
                img = _render_formula(formula_text, output_dir, formula_idx[0])
                if img:
                    doc.add_picture(img, width=Inches(4))
                formula_idx[0] += 1
            continue

        # ── 行内公式 $...$ ──
        if "$" in line:
            # 简单处理：渲染出公式图片，替换文本
            inline_formulas = re.findall(r"\$(.*?)\$", line)
            for f in inline_formulas:
                img = _render_formula(f, output_dir, formula_idx[0])
                if img:
                    doc.add_picture(img, width=Inches(2))
                formula_idx[0] += 1
            # 去掉公式部分后把剩余文本作为段落
            rest = re.sub(r"\$.*?\$", "", line).strip()
            if rest:
                p = doc.add_paragraph()
                _add_run(p, rest)
            i += 1
            continue

        # ── 表格 ──
        if "|" in line and line.strip().startswith("|"):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                stripped = lines[i].strip()
                # 跳过分隔行（|---|）
                if re.match(r"^\|[\s\-:]+\|$", stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.columns):
                            table.rows[ri].cells[ci].text = cell
            continue

        # ── 引用块 ──
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(quote_lines))
            run.italic = True
            continue

        # ── 标题 ──
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            doc.add_heading(title, level=level)
            i += 1
            continue

        # ── 列表 ──
        ul_match = re.match(r"^[\s]*[-*+]\s+(.+)$", line)
        ol_match = re.match(r"^[\s]*\d+\.\s+(.+)$", line)
        if ul_match or ol_match:
            items = []
            is_ordered = ol_match is not None
            while i < len(lines):
                m = re.match(r"^[\s]*[-*+]\s+(.+)$", lines[i])
                om = re.match(r"^[\s]*\d+\.\s+(.+)$", lines[i])
                if m or om:
                    items.append((m or om).group(1))
                    i += 1
                else:
                    break
            for item in items:
                p = doc.add_paragraph(style="List Bullet" if not is_ordered else "List Number")
                _add_run(p, item)
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        _add_run(p, line.strip())
        i += 1

    # ── 保存 ──
    if not output_path:
        output_path = os.path.join(output_dir, f"output_{uuid.uuid4().hex[:8]}.docx")
    doc.save(output_path)
    return os.path.abspath(output_path)
