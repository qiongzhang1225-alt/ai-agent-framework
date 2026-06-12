"""文件操作工具（Claude Code 风格的精细化代码助手能力）：

- ``read_file``  按扩展名自动解析（text/xlsx/pdf/docx）
- ``write_file`` 默认版本化（不覆盖原文件）
- ``edit_file``  精确字符串替换（结果写新版本）
- ``grep``       正则搜索
- ``glob``       glob 列文件
"""
from __future__ import annotations

import re
from pathlib import Path

from ai_agent import tool
from paths import DEFAULT_WORKDIR
from tools._common import safe_workdir_path


# ── 共享常量 ─────────────────────────────────────────────────────────────────

# 文本文件扩展名（grep / read_file 用来判断是否按文本处理）
_TEXT_EXTS = frozenset({
    ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".log", ".csv", ".tsv",
    ".js", ".ts", ".jsx", ".tsx", ".html", ".htm", ".css", ".scss",
    ".xml", ".ini", ".toml", ".cfg", ".conf", ".bat", ".sh", ".ps1",
    ".java", ".c", ".cpp", ".h", ".hpp", ".go", ".rs", ".rb", ".php",
    ".sql", ".env", ".gitignore", ".editorconfig",
})

# 搜索 / 列文件时跳过的"噪音"目录
_SKIP_DIRS = frozenset({
    "__pycache__", ".git", ".venv", "venv", "env", "myenv", "virtualenv",
    "node_modules",
    ".sandbox", ".memory", ".idea", ".vscode", ".pytest_cache",
    ".mypy_cache", "dist", "build", ".next",
    "site-packages",  # pip install 目标，路径里任何层级出现都跳过
})


def _is_venv_dir(path: Path) -> bool:
    """识别 Python 虚拟环境根目录（有 pyvenv.cfg 文件即是）。

    用 pyvenv.cfg 当 ground truth 比硬编码"".venv""venv""env" 名单更可靠
    —— 用户用什么名字建 venv 都不漏（myenv / dev / test_env 等）。
    """
    try:
        return (path / "pyvenv.cfg").is_file()
    except Exception:
        return False


def _is_pkg_metadata_dir(name: str) -> bool:
    """识别 pip 包元数据目录（*.egg-info / *.dist-info）。"""
    return name.endswith(".egg-info") or name.endswith(".dist-info")

_MAX_FILE_CHARS = 8000   # read_file 输出截断阈值
_MAX_GREP_MATCHES = 50   # grep 返回最多条数
_MAX_GLOB_MATCHES = 100  # glob 返回最多条数


def _slice_text_by_lines(text: str, offset: int, limit: int, max_chars: int) -> str:
    """按行切片文本，输出带行号（``  123\\t内容``）。

    供 ``read_file`` / ``self_read_file`` 共用。

    Args:
        text: 原文
        offset: 1-indexed 起始行（0 视作 1）
        limit: 最多读多少行（0 = 不限）
        max_chars: 输出字符上限（超出再截断）

    返回带 "(显示第 N-M 行，共 K 行)" 头部 + cat -n 风格内容。
    """
    lines = text.split("\n")
    total = len(lines)
    start = max(0, offset - 1) if offset > 0 else 0
    if start >= total:
        return f"(offset={offset} 超过文件总行数 {total})"
    end = start + limit if limit > 0 else total
    end = min(end, total)
    sliced = lines[start:end]
    width = max(len(str(end)), 3)
    rendered = "\n".join(f"{i + start + 1:>{width}}\t{ln}" for i, ln in enumerate(sliced))
    header = f"(显示第 {start + 1}-{end} 行，共 {total} 行)\n"
    if len(rendered) > max_chars:
        rendered = rendered[:max_chars] + "\n...(单段超长，请调小 limit)"
    return header + (rendered or "(空切片)")


def _next_version_path(target: Path) -> Path:
    """根据 target 文件名找下一个可用的版本号路径。

    规则：
    - ``report.txt`` 已存在 → 找 ``report_v2.txt``，存在再试 ``v3``、``v4``...
    - ``report_v3.txt`` → 视作 v3，从 v4 开始找
    - 文件名末尾的 ``_v<数字>`` 模式会被识别为版本号

    返回的路径**保证不存在**（可以直接写入）。
    """
    stem = target.stem
    ext = target.suffix
    parent = target.parent

    m = re.match(r"^(.*)_v(\d+)$", stem)
    if m:
        base = m.group(1)
        start_v = int(m.group(2)) + 1
    else:
        base = stem
        start_v = 2  # 原始视为 v1，新版本从 v2 起

    n = start_v
    while (parent / f"{base}_v{n}{ext}").exists():
        n += 1
    return parent / f"{base}_v{n}{ext}"


def _backup_target(target: Path) -> Path:
    """累积式备份：先把 target 复制成 ``<file>.bak`` / ``.bak2`` / ``.bak3`` ...
    找下一个未占用的 .bak 后缀写入，返回 backup 路径。

    供 write_file(force=True) 和 edit_file 共用 —— 任何"覆盖原文件"行为前
    都先 backup，C2 安全策略要求。
    """
    bak_root = target.with_suffix(target.suffix + ".bak")
    if not bak_root.exists():
        bak_path = bak_root
    else:
        n = 2
        while True:
            candidate = target.parent / f"{bak_root.name}{n}"
            if not candidate.exists():
                bak_path = candidate
                break
            n += 1
    bak_path.write_bytes(target.read_bytes())
    return bak_path


# ── read_file ───────────────────────────────────────────────────────────────

@tool
def read_file(path: str, config: dict, offset: int = 0, limit: int = 0) -> str:
    """读取工作目录中的文件，按扩展名自动解析。

    支持格式：
    - 文本（.txt / .md / .py / .json / .yaml / .csv 等）：按 UTF-8 读
    - Excel（.xlsx / .xls）：pandas 读，每个 sheet 前 50 行
    - PDF（.pdf）：pdfplumber 提取每页文本
    - Word（.docx）：python-docx 提取段落 + 表格文本
    - 其他二进制：仅返回文件信息提示

    所有路径必须在当前会话的工作目录内（绝对路径必须在 workdir 内子路径）。
    超过 8000 字符的内容会被截断（防止 token 浪费）。

    **行切片**（offset / limit）：
    - 默认 ``offset=0, limit=0`` → 整文件读（保持原行为）
    - 指定后只读指定行段，输出带行号（``  123\\t内容``），方便后续 edit_file 引用
    - 仅对**文本类**文件生效；Excel/PDF/Word 走自己的分页逻辑，忽略本参数
    - 看 2000 行的大文件时，建议先整读拿到总行数提示，再按段 offset/limit 精读

    Args:
        path: 文件路径（相对路径相对 workdir，或工作目录内的绝对路径）
        offset: 从第几行开始读（1-indexed；0 = 从头）
        limit: 最多读多少行（0 = 不限，仍受 8000 字符上限约束）
    """
    try:
        target = safe_workdir_path(path, config, must_exist=True)
    except ValueError as e:
        return f"读取失败: {e}"

    if not target.is_file():
        return f"读取失败: {target} 不是文件"

    ext = target.suffix.lower()

    try:
        if ext in _TEXT_EXTS or ext == "":
            text = target.read_text(encoding="utf-8", errors="replace")
            # 行切片分支（offset / limit 至少一个 > 0）
            if offset > 0 or limit > 0:
                return _slice_text_by_lines(text, offset, limit, _MAX_FILE_CHARS)
            # 默认整文件读
            if len(text) > _MAX_FILE_CHARS:
                total_lines = text.count("\n") + 1
                return (
                    text[:_MAX_FILE_CHARS]
                    + f"\n\n...(已截断，原文 {len(text)} 字符 / {total_lines} 行)"
                    f"\n提示: 用 read_file(path, offset=N, limit=M) 按行精读后续部分"
                )
            return text or "(空文件)"

        if ext in (".xlsx", ".xls"):
            import pandas as pd
            xls = pd.ExcelFile(target)
            parts = []
            for sheet in xls.sheet_names:
                df = pd.read_excel(target, sheet_name=sheet, nrows=50)
                parts.append(f"=== Sheet: {sheet} (前 50 行, 共 {df.shape[0]} 行) ===\n{df.to_string()}")
            text = "\n\n".join(parts)
            return text[:_MAX_FILE_CHARS] + "\n...(已截断)" if len(text) > _MAX_FILE_CHARS else text

        if ext == ".pdf":
            import pdfplumber
            parts = []
            with pdfplumber.open(target) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text() or "(无文字内容)"
                    parts.append(f"=== Page {i} ===\n{t}")
            text = "\n\n".join(parts)
            return text[:_MAX_FILE_CHARS] + "\n...(已截断)" if len(text) > _MAX_FILE_CHARS else text

        if ext == ".docx":
            from docx import Document
            doc = Document(str(target))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
            return text[:_MAX_FILE_CHARS] + "\n...(已截断)" if len(text) > _MAX_FILE_CHARS else text or "(空文档)"

        size = target.stat().st_size
        return f"(二进制文件 {ext}, {size} 字节, 无法以文本形式预览。如需处理请用 execute_code)"
    except Exception as e:
        return f"读取失败 ({type(e).__name__}): {e}"


# ── write_file ──────────────────────────────────────────────────────────────

@tool
def write_file(path: str, content: str, config: dict, force: bool = False) -> str:
    """写入文件到工作目录。**默认版本化**，不覆盖已有文件。

    行为规则：
    - 目标不存在 → 直接创建
    - 目标已存在 + ``force=False``（默认）→ 写到 ``<stem>_v2.<ext>``（自动避开
      已占用的版本号），**原文件不动**
    - 目标已存在 + ``force=True`` → 真正覆盖原文件，先备份为
      ``<path>.bak`` 兜底（仅首次备份）

    什么时候用 ``force=True``：
    - 用户**明确**说"直接覆盖 / 别留版本号 / 替换原文件"
    - 用户没明确表态时一律 ``force=False``（默认）

    UTF-8 编码；父目录自动创建。

    Args:
        path: 文件路径
        content: 要写入的全文
        force: 强制覆盖原文件（默认 False，启用版本化保护）
    """
    try:
        target = safe_workdir_path(path, config, must_exist=False)
    except ValueError as e:
        return f"写入失败: {e}"

    target.parent.mkdir(parents=True, exist_ok=True)

    # 目标不存在 → 直接写
    if not target.exists():
        target.write_text(content, encoding="utf-8")
        return f"已写入 {target.name} ({target.stat().st_size} 字节)"

    # 目标已存在 + 非强制 → 版本化
    if not force:
        new_target = _next_version_path(target)
        new_target.write_text(content, encoding="utf-8")
        return (
            f"已生成新版本 {new_target.name} ({new_target.stat().st_size} 字节)；"
            f"原文件 {target.name} 未改动。"
            "（如需直接覆盖原文件，调用时传 force=True）"
        )

    # 强制覆盖 → 累积式 .bak 兜底再覆盖（C2 保护）
    bak_path = _backup_target(target)
    target.write_text(content, encoding="utf-8")
    return (
        f"已覆盖 {target.name} ({target.stat().st_size} 字节)；"
        f"原内容备份为 {bak_path.name}"
    )


# ── edit_file ───────────────────────────────────────────────────────────────

@tool
def edit_file(path: str, old_string: str, new_string: str, config: dict) -> str:
    """对工作目录内文件做**精确字符串替换**，**直接修改原文件**（先累积 .bak 备份）。

    行为（已和 Claude Code 的 Edit 对齐）：
    - 读源文件 → 找到 old_string 的位置 → 替换 → **直接写回源文件**
    - 改前自动备份到 ``<file>.bak`` / ``.bak2`` / ``.bak3`` ...（绝不覆盖旧备份）
    - 改错了？让主人告诉你，你可以 read_file 看 ``<file>.bak`` 找回原文

    严格匹配规则：
    - old_string 必须在源文件中**唯一出现**（0 次或多次都会失败）
    - 不唯一时请扩展 old_string 加入更多上下文（缩进 / 空格 / 换行必须精确匹配）

    什么时候**不要**用 edit_file：
    - 整文件重写 → 用 write_file(path, content, force=True)，更直接
    - 多处分散修改 → 一次次 edit_file 调用 OK，但每次都会生成新 .bak

    Args:
        path: 源文件路径
        old_string: 要替换的原文（必须唯一）
        new_string: 替换成的新文本
    """
    try:
        source = safe_workdir_path(path, config, must_exist=True)
    except ValueError as e:
        return f"编辑失败: {e}"

    if not source.is_file():
        return f"编辑失败: {source} 不是文件"

    try:
        text = source.read_text(encoding="utf-8")
    except Exception as e:
        return f"编辑失败: 无法读取（可能是二进制文件）: {e}"

    if old_string == new_string:
        return "编辑失败: old_string 和 new_string 相同，无操作"

    count = text.count(old_string)
    if count == 0:
        return (
            f"编辑失败: old_string 未在 {source.name} 中出现。"
            "请检查空白字符 / 缩进 / 大小写是否完全匹配。"
        )
    if count > 1:
        return (
            f"编辑失败: old_string 在 {source.name} 中出现 {count} 次，必须唯一。"
            "请扩展 old_string 加入更多上下文使其唯一。"
        )

    # 改前累积 .bak 兜底（C2 保护）
    bak_path = _backup_target(source)
    new_text = text.replace(old_string, new_string, 1)
    source.write_text(new_text, encoding="utf-8")

    return (
        f"已编辑 {source.name}（替换 1 处；"
        f"原内容备份为 {bak_path.name}）"
    )


# ── grep ────────────────────────────────────────────────────────────────────

@tool
def grep(pattern: str, config: dict, path: str = ".") -> str:
    """在工作目录内按正则搜索文件内容。

    - pattern: Python 正则表达式（如 'def \\w+'、'TODO|FIXME'）
    - path: 搜索范围；默认 "." 搜整个 workdir；可指定子目录或单个文件
    - 自动跳过 __pycache__ / .git / .venv / node_modules 等噪音目录
    - 仅搜索文本文件（按扩展名白名单）
    - 返回最多 50 个匹配，格式 "<相对路径>:<行号>: <匹配的整行>"

    Args:
        pattern: 正则
        path: 搜索起点（默认整个工作目录）
    """
    cfg = (config or {}).get("configurable", {}) if config else {}
    workdir = Path(cfg.get("workdir") or str(DEFAULT_WORKDIR)).resolve()

    try:
        target = safe_workdir_path(path, config, must_exist=True)
    except ValueError as e:
        return f"grep 失败: {e}"

    try:
        regex = re.compile(pattern)
    except re.error as e:
        return f"grep 失败: 无效的正则 {pattern!r}: {e}"

    matches: list[str] = []

    def search_one(fp: Path) -> bool:
        """搜索单个文件；返回 False 表示已达到上限应中止外层循环。"""
        try:
            with fp.open("r", encoding="utf-8", errors="replace") as f:
                for lineno, line in enumerate(f, 1):
                    if regex.search(line):
                        rel = fp.relative_to(workdir).as_posix()
                        matches.append(f"{rel}:{lineno}: {line.rstrip()}")
                        if len(matches) >= _MAX_GREP_MATCHES:
                            return False
        except Exception:
            pass
        return True

    if target.is_file():
        search_one(target)
    else:
        for fp in target.rglob("*"):
            if not fp.is_file():
                continue
            try:
                rel_parts = fp.relative_to(workdir).parts
            except ValueError:
                continue
            if any(part in _SKIP_DIRS for part in rel_parts):
                continue
            if fp.suffix.lower() not in _TEXT_EXTS and fp.suffix != "":
                continue
            if not search_one(fp):
                break

    if not matches:
        return f"无匹配（pattern={pattern!r}, path={path!r}）"
    result = "\n".join(matches)
    if len(matches) >= _MAX_GREP_MATCHES:
        result += f"\n...(达到 {_MAX_GREP_MATCHES} 条上限)"
    return result


# ── glob ────────────────────────────────────────────────────────────────────

@tool
def glob(pattern: str, config: dict) -> str:
    """按 glob 模式列出工作目录内的文件。

    模式语法：
    - 单 ``*`` 匹配同级任意非斜杠字符
    - ``**`` 匹配任意层级（含 0 层）
    - ``?`` / ``[...]`` 单字符 / 字符集
    - 例: ``*.py``、``**/*.xlsx``、``src/**/*.ts``

    自动跳过 __pycache__ / .git / .venv 等。返回最多 100 个文件，
    路径按字典序排序，相对工作目录显示。

    Args:
        pattern: glob 模式
    """
    cfg = (config or {}).get("configurable", {}) if config else {}
    workdir = Path(cfg.get("workdir") or str(DEFAULT_WORKDIR)).resolve()
    workdir.mkdir(parents=True, exist_ok=True)

    try:
        results: list[str] = []
        for fp in workdir.glob(pattern):
            if not fp.is_file():
                continue
            try:
                rel_parts = fp.relative_to(workdir).parts
            except ValueError:
                continue
            if any(part in _SKIP_DIRS for part in rel_parts):
                continue
            results.append(fp.relative_to(workdir).as_posix())
            if len(results) >= _MAX_GLOB_MATCHES:
                break
        results.sort()
    except Exception as e:
        return f"glob 失败 ({type(e).__name__}): {e}"

    if not results:
        return f"无匹配（pattern={pattern!r}）"
    out = "\n".join(results)
    if len(results) >= _MAX_GLOB_MATCHES:
        out += f"\n...(达到 {_MAX_GLOB_MATCHES} 条上限)"
    return out
