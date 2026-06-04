"""FastAPI 后端 —— 替代 Streamlit，提供自由可定制的 UI。

后端职责：
- 管理对话（内存中的会话状态）
- 工作目录 CRUD
- SSE 流式聊天接口
- 文件上传 / 下载 / 预览
- 工作目录文件列表
"""

from __future__ import annotations

import asyncio
import json
import mimetypes
import os
import re
import sys
import uuid
from pathlib import Path
from typing import Any, AsyncIterator

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse

from agent import create_agent
from ai_agent import (
    JSONCheckpoint, Message, message_from_dict, message_to_dict,
)
from paths import (
    PROJECT_ROOT,
    ASSETS_DIR,
    STATIC_DIR,
    TEMPLATES_DIR,
)

# ── 配置 ─────────────────────────────────────────────────────────────────────
# 数据目录跟 PROJECT_ROOT（打包模式下 = exe 旁，源码模式下 = 项目根）
# 资源目录（assets/static/templates）由 paths.py 处理：
#   - 源码模式：PROJECT_ROOT/{assets,static,templates}
#   - 打包模式：sys._MEIPASS/{static,templates}，assets 优先 exe 旁
WORKSPACE_ROOT = PROJECT_ROOT / ".sandbox" / "workspace"
META_ROOT = PROJECT_ROOT / ".sandbox" / "_meta"

MODELS: dict[str, str] = {
    "deepseek-v4-flash": "DeepSeek V4 Flash",
    "deepseek-v4-pro": "DeepSeek V4 Pro",
    "mimo-v2.5": "MiMo v2.5（视觉）",
    "mimo-v2.5-pro": "MiMo v2.5 Pro",
}
DEFAULT_MODEL = "deepseek-v4-flash"
# 注：视觉能力通过 vision_describe 工具调 MiMo 实现（DeepSeek 主导 + MiMo 识别），
# 不再有"自动切换模型"逻辑（详见 tools/vision.py 注释）。

# ── App ──────────────────────────────────────────────────────────────────────

app = FastAPI(title="信息统合思念体")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.mount("/assets", StaticFiles(directory=str(ASSETS_DIR)), name="assets")


# 启动时打一次全量快照（C3 保护策略）。同一天已快照过会自动跳过，
# 失败永不阻塞启动。这是给"删错了发现得晚"留的远期保险。
@app.on_event("startup")
async def _on_startup_backup():
    try:
        from backups import snapshot_all
        summary = snapshot_all()
        msgs = []
        if summary["memory_snapshot"]:
            msgs.append(f"memory→{summary['memory_snapshot']}")
        if summary["skills_snapshot"]:
            msgs.append(f"skills→{summary['skills_snapshot']}")
        if summary["memory_cleaned"] or summary["skills_cleaned"]:
            msgs.append(f"清理过期 memory={summary['memory_cleaned']} skills={summary['skills_cleaned']}")
        if msgs:
            print(f"[backups] 启动快照: {' | '.join(msgs)}")
    except Exception as e:
        print(f"[backups] 启动快照失败（不影响主服务）: {e}")


# 启动时兜底 commit working tree（防 self_edit 工具被改坏跳过 commit）。
# 即便有希改了 tools/self_edit.py 让"以后不 commit"，本 hook 仍会兜底落账。
# 失败永不阻塞启动。
@app.on_event("startup")
async def _on_startup_auto_commit():
    try:
        from tools.self_edit import auto_commit_pending
        h = auto_commit_pending("server-startup")
        if h:
            print(f"[self-edit] 启动时兜底 commit: {h}")
    except Exception as e:
        print(f"[self-edit] 启动兜底 commit 失败（不影响主服务）: {e}")

# ── 全局状态 ─────────────────────────────────────────────────────────────────

# Phase 3 之后：LangGraph 的 SqliteSaver 不再使用；对话上下文完整存进
# conv.json（含 tool_calls / reasoning_content），重启后由 _load_into_globals
# 加载，再由 chat 路由从 conv["messages"] 重建 list[Message] 喂给 Agent。

_agents: dict[str, Any] = {}  # 兼容遗留代码（不再被 chat 路由用）


def get_agent(model_id: str):
    """**遗留**接口（兼容）：按 model 缓存 Agent。

    新代码应该用 ``get_agent_for_conv(conv)`` —— 它会根据 conv 类型动态拼装
    prompt（master 会注入子对话摘要）。这个旧接口仅供尚未迁移的代码使用。
    """
    if model_id not in _agents:
        _agents[model_id] = create_agent(model_id)
    return _agents[model_id]


def get_agent_for_conv(conv: dict):
    """根据 conv 动态创建 Agent。

    每次 chat 都新建（不缓存 Agent 实例），因为 prompt 会根据 conv 类型
    变化（master 注入子摘要、新批准的子摘要要被看到等）。LLM client 由
    ``create_agent`` 内部各自创建 —— httpx 异步客户端轻量，每次创建~10ms。

    若未来要优化可以加 LLM cache，但实测当前性能足够。
    """
    return create_agent(conv.get("model", DEFAULT_MODEL), conv=conv)


# 会话：{thread_id: {id, name, messages, workdir, model}}
conversations: dict[str, dict] = {}


# Phase 4: 用 ai_agent.JSONCheckpoint 替换原来散在 server.py 里的持久化函数
_checkpoint = JSONCheckpoint(META_ROOT, filename="conv.json")


def _load_into_globals() -> None:
    """启动时调用：把所有持久化的对话加载进内存 `conversations` 字典。"""
    for tid, data in _checkpoint.load_all().items():
        # 校验 workdir 仍存在；如不存在尝试重建（用户可能改过路径）
        if data.get("workdir"):
            try:
                Path(data["workdir"]).mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
        # 旧 conv 兼容：补缺失字段（kind / parent_id / sub_level / summary 等）
        conversations[tid] = _ensure_conv_fields(data)


MASTER_CONV_ID = "master_yuki"
MASTER_CONV_NAME = "有希"


def new_conversation(
    model_id: str = DEFAULT_MODEL,
    *,
    kind: str = "standalone",
    parent_id: str | None = None,
    sub_level: str | None = None,
    fixed_id: str | None = None,
    name: str | None = None,
) -> dict:
    """创建对话。``kind`` 可选：

    - ``"standalone"``（默认，独立 thread，老行为）
    - ``"master"``（主对话，全局只有一个，id 固定为 MASTER_CONV_ID）
    - ``"sub"``（子对话，挂在某个 master 下，必须传 parent_id 和 sub_level）

    sub_level 仅 sub 有意义：``"advanced"`` 或 ``"restricted"``。
    """
    tid = fixed_id or (MASTER_CONV_ID if kind == "master" else str(uuid.uuid4()))
    workdir = WORKSPACE_ROOT / tid
    workdir.mkdir(parents=True, exist_ok=True)
    conv: dict[str, Any] = {
        "id": tid,
        "name": name or (MASTER_CONV_NAME if kind == "master" else "新对话"),
        "messages": [],
        "workdir": str(workdir),
        "model": model_id,
        # 新增字段（M+ 设计）
        "kind": kind,
        "parent_id": parent_id,
        "sub_level": sub_level if kind == "sub" else None,
        "summary": "",
        "summary_updated_at": "",
        "summary_pending_approval": False,
        "summary_approved_for_master": False,
        "last_compress_at": "",
    }
    conversations[tid] = conv
    _checkpoint.save(conv["id"], conv)
    return conv


def ensure_master_conversation() -> dict:
    """启动时确保 master 对话存在；不存在则创建。返回 master conv。"""
    existing = conversations.get(MASTER_CONV_ID)
    if existing and existing.get("kind") == "master":
        return existing
    # 找任何 kind==master 的（防 id 改名等）
    for c in conversations.values():
        if c.get("kind") == "master":
            return c
    return new_conversation(kind="master")


def _ensure_conv_fields(conv: dict) -> dict:
    """给旧 conv 补缺失字段，向后兼容。**就地修改**。"""
    conv.setdefault("kind", "standalone")
    conv.setdefault("parent_id", None)
    conv.setdefault("sub_level", None)
    conv.setdefault("summary", "")
    conv.setdefault("summary_updated_at", "")
    conv.setdefault("summary_pending_approval", False)
    conv.setdefault("summary_approved_for_master", False)
    conv.setdefault("last_compress_at", "")
    conv.setdefault("compressed_summaries", [])  # 每次压缩 append 一条 {ts, first_ts, last_ts, count, memory_id, summary_preview}
    return conv


# ── 启动序列：加载持久化对话 + 确保 master 存在 ────────────────────────────
_load_into_globals()
ensure_master_conversation()


# ── 后台预热代码索引（不阻塞启动）────────────────────────────────────────
# yuki 调 code_search / code_outline 等工具时不用等数秒首次索引
# 失败不致命（tree-sitter 解析异常等），yuki 首次调时还会触发
def _warmup_code_index():
    try:
        import threading
        from tools.code_indexer import get_indexer

        def _do_warmup():
            try:
                get_indexer().refresh(str(PROJECT_ROOT))
            except Exception as e:
                print(f"[code_indexer 预热失败（不影响功能）] {e}")

        threading.Thread(
            target=_do_warmup, daemon=True, name="code_index_warmup"
        ).start()
    except Exception as e:
        print(f"[code_indexer 启动跳过] {e}")


_warmup_code_index()


# ── 工作目录文件过滤（防 venv / node_modules / 缓存等淹没列表）────────────
# 目录黑名单：遇到这些名字的子目录整个跳过递归
_EXCLUDE_DIRS = frozenset({
    ".venv", "venv", "env", ".env_dir",
    "node_modules", "__pycache__",
    ".git", ".idea", ".vscode",
    ".pytest_cache", ".mypy_cache", ".ruff_cache", ".tox",
    "dist", "build", "target",
    ".next", ".nuxt", ".svelte-kit", ".turbo",
    "_internal",                  # PyInstaller onedir 产物
    ".execute_trash", ".skills_trash",
    ".memory", ".memory_backups", ".skills_backups",
    ".sandbox",                   # 内部对话状态
    "models",                     # 本地模型（动辄几百 MB）
})

# 文件后缀黑名单
_EXCLUDE_SUFFIXES = frozenset({
    ".pyc", ".pyo", ".pyd",
    ".class", ".o", ".obj",
    ".tmp", ".lock", ".swp", ".swo",
    ".log",                       # 防止日志狂刷
})

# 完整文件名黑名单
_EXCLUDE_FILES = frozenset({
    ".DS_Store", "Thumbs.db", "Desktop.ini",
    ".yuki.lock", ".yuki-launcher.log",
})


def _file_excluded(path: Path) -> bool:
    """单文件级别过滤（已经知道是 file 才调，dir 在 walk 时另判）。"""
    name = path.name
    if name in _EXCLUDE_FILES:
        return True
    if path.suffix.lower() in _EXCLUDE_SUFFIXES:
        return True
    # *.bak / *.bak2 / *.bak3 ... 累积备份
    if name.endswith(".bak") or re.search(r"\.bak\d+$", name):
        return True
    return False


def _walk_workdir(workdir: Path):
    """生成器：遍历 workdir，跳过 _EXCLUDE_DIRS，过滤垃圾文件。

    yield 元组 (rel_path: str, abs_path: Path, is_dir: bool)。
    """
    if not workdir.exists():
        return
    workdir = workdir.resolve()
    # 用 os.walk 比 Path.rglob 快且能截断递归
    for root, dirs, files in os.walk(workdir):
        # 原地修改 dirs 跳过黑名单目录（os.walk 这是 idiom）
        dirs[:] = [d for d in dirs if d not in _EXCLUDE_DIRS]
        root_path = Path(root)
        rel_root = root_path.relative_to(workdir)
        # 目录本身（除 workdir 根）
        if rel_root != Path("."):
            yield (str(rel_root).replace("\\", "/"), root_path, True)
        for fname in files:
            fpath = root_path / fname
            if _file_excluded(fpath):
                continue
            rel = (rel_root / fname).as_posix() if rel_root != Path(".") else fname
            yield (rel, fpath, False)


def _list_workdir_files(workdir: str) -> set[Path]:
    """返回 workdir 下所有 file 的 Path set（已过滤黑名单）。

    用于 new_files diff（done 事件计算 yuki 本轮生成的文件）。
    """
    root = Path(workdir)
    return {abs_p for _, abs_p, is_dir in _walk_workdir(root) if not is_dir}


def _safe_path(workdir: str, relpath: str) -> Path:
    """把相对路径解析为绝对路径，并校验落在 workdir 内（防 path traversal）。"""
    workdir_abs = Path(workdir).resolve()
    target = (workdir_abs / relpath).resolve()
    try:
        target.relative_to(workdir_abs)
    except ValueError:
        raise HTTPException(403, "Path escapes workdir")
    return target


# ── 首页 ─────────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    if not conversations:
        new_conversation()
    has_bg = any(
        (ASSETS_DIR / f"background.{ext}").exists()
        for ext in ("png", "jpg", "jpeg", "webp")
    )
    bg_url = None
    if has_bg:
        for ext in ("png", "jpg", "jpeg", "webp"):
            if (ASSETS_DIR / f"background.{ext}").exists():
                bg_url = f"/assets/background.{ext}"
                break
    has_icon = (ASSETS_DIR / "icon.png").exists()
    return templates.TemplateResponse(
        request,
        "index.html",
        {
            "models": MODELS,
            "default_model": DEFAULT_MODEL,
            "background_url": bg_url,
            "icon_url": "/assets/icon.png" if has_icon else None,
        },
    )


# ── 会话 CRUD ────────────────────────────────────────────────────────────────

def _conv_summary(conv: dict) -> dict:
    """会话列表的精简表示。包含 sidebar 渲染所需的所有元数据。"""
    msgs = conv.get("messages") or []
    last_ts = ""
    for m in reversed(msgs):
        if m.get("ts"):
            last_ts = m["ts"]
            break

    # 兜底（关键 fix）：消息没 ts 时（新对话 / N1 之前的旧对话）用
    # conv.json 文件的 mtime 作为"最近活跃时间"。这样：
    # - 刚新建的对话进"一天前"桶（mtime ≈ 现在）
    # - 旧对话进自己实际修改时间对应的桶
    if not last_ts:
        try:
            from datetime import datetime as _dt
            conv_path = _checkpoint._path(conv["id"])
            if conv_path.exists():
                last_ts = _dt.fromtimestamp(conv_path.stat().st_mtime).isoformat(timespec="seconds")
        except Exception:
            pass

    return {
        "id": conv["id"],
        "name": conv["name"],
        "workdir": conv["workdir"],
        "model": conv["model"],
        "message_count": len(msgs),
        # 主-子设计字段（前端 sidebar / 徽章 / 抽屉分组用）
        "kind": conv.get("kind", "standalone"),
        "parent_id": conv.get("parent_id"),
        "sub_level": conv.get("sub_level"),
        "summary_pending_approval": conv.get("summary_pending_approval", False),
        "summary_approved_for_master": conv.get("summary_approved_for_master", False),
        "last_message_ts": last_ts,
        # 前端 messages 数组：只暴露 ts 给时间桶排序用，不暴露内容（list 接口要轻）
        "messages": [{"ts": last_ts}] if last_ts else [],
    }


@app.get("/api/health")
async def health() -> dict:
    """轻量探活端点 —— launcher 的 splash 画面用来判断 server 是否就绪。

    可选 ``?warmed=1`` 参数：要求 chromadb embedding 模型已加载完
    （懒加载首次访问要 3-5 秒）。launcher 会先调一次预热，再切主界面。
    """
    from fastapi import Request
    return {"ok": True, "convs": len(conversations)}


@app.get("/api/conversations")
async def list_conversations() -> list[dict]:
    # 倒序：最新创建的在前
    return [_conv_summary(c) for c in reversed(list(conversations.values()))]


@app.get("/api/conversations/{tid}")
async def get_conversation(tid: str) -> dict:
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    return conversations[tid]


class CreateConvRequest(BaseModel):
    model: str = DEFAULT_MODEL
    kind: str = "standalone"          # "standalone" / "sub"（master 由系统自建）
    parent_id: str | None = None      # sub 时必填
    sub_level: str | None = None      # sub 时必填：advanced / restricted


@app.post("/api/conversations")
async def create_conversation(req: CreateConvRequest) -> dict:
    if req.model not in MODELS:
        raise HTTPException(400, "Unknown model")
    # master 不能通过 API 创建（系统启动时自建唯一一个）
    if req.kind not in ("standalone", "sub"):
        raise HTTPException(400, f"kind 只能是 standalone 或 sub，不接受 {req.kind!r}")
    if req.kind == "sub":
        if not req.parent_id or req.parent_id not in conversations:
            raise HTTPException(400, "sub 必须传有效的 parent_id")
        if conversations[req.parent_id].get("kind") != "master":
            raise HTTPException(400, "parent 必须是 master 对话")
        if req.sub_level not in ("advanced", "restricted"):
            raise HTTPException(400, "sub_level 必须是 advanced 或 restricted")
    return new_conversation(
        req.model,
        kind=req.kind,
        parent_id=req.parent_id,
        sub_level=req.sub_level,
    )


@app.delete("/api/conversations/{tid}")
async def delete_conversation(tid: str):
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    # master 不能删（防误删唯一主对话）
    if conversations[tid].get("kind") == "master":
        raise HTTPException(400, "主对话不可删除")
    del conversations[tid]
    _checkpoint.delete(tid)
    # 保证至少有一个对话存在（master 兜底已存在所以这条几乎不会触发）
    if not conversations:
        new_conversation()
        ensure_master_conversation()
    return {"ok": True}


class UpdateConvRequest(BaseModel):
    name: str | None = None
    workdir: str | None = None
    model: str | None = None
    sub_level: str | None = None      # 子对话权限切换（advanced / restricted）


@app.patch("/api/conversations/{tid}")
async def update_conversation(tid: str, req: UpdateConvRequest) -> dict:
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    conv = conversations[tid]
    if req.name is not None:
        # master 不能改名（防止用户误改后找不到"有希"）
        if conv.get("kind") == "master":
            raise HTTPException(400, "主对话名字固定，不能改")
        conv["name"] = req.name.strip() or "未命名"
    if req.sub_level is not None:
        if conv.get("kind") != "sub":
            raise HTTPException(400, "只有子对话可改 sub_level")
        if req.sub_level not in ("advanced", "restricted"):
            raise HTTPException(400, "sub_level 必须是 advanced 或 restricted")
        conv["sub_level"] = req.sub_level
    if req.workdir is not None:
        try:
            path = Path(req.workdir).expanduser().resolve()
            path.mkdir(parents=True, exist_ok=True)
            conv["workdir"] = str(path)
        except Exception as e:
            raise HTTPException(400, f"Invalid workdir: {e}")
    if req.model is not None:
        if req.model not in MODELS:
            raise HTTPException(400, "Unknown model")
        conv["model"] = req.model
    _checkpoint.save(conv["id"], conv)
    return conv


# ── SSE 聊天 ─────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    thread_id: str
    message: str = ""           # resume=True 时可空
    # 多模态：可选的图片 base64 data URL 列表（"data:image/jpeg;base64,..."）
    # 非空时 server 会把消息 content 拼成 OpenAI 多模态 list 格式
    images: list[str] = []
    # 断点续传：上一轮撞 max_iterations 后，前端点"继续"按钮触发；
    # True 时**不 append user message**，直接用现有 history 让 LLM 接着跑
    resume: bool = False


def _save_image_for_conv(tid: str, data_url: str) -> str | None:
    """把前端传来的 base64 data URL 存到 .sandbox/_meta/<tid>/images/。

    返回 image_id（含 ``img_`` 前缀，对应 vision_describe 工具的入参）。
    失败返回 None（不抛 —— 不该让一张坏图毁掉整轮对话）。
    """
    if not data_url or not data_url.startswith("data:"):
        return None
    try:
        # data:image/png;base64,iVBORw0K...
        head, _, b64 = data_url.partition(",")
        if not b64:
            return None
        # mime 在 "data:image/xxx;base64" 部分
        mime = head.split(";")[0].removeprefix("data:") or "image/png"
        ext = mimetypes.guess_extension(mime) or ".png"
        # 防御性：奇怪的扩展名（.jpe → .jpg）
        if ext == ".jpe":
            ext = ".jpg"

        img_dir = META_ROOT / tid / "images"
        img_dir.mkdir(parents=True, exist_ok=True)
        image_id = f"img_{uuid.uuid4().hex[:8]}"
        path = img_dir / f"{image_id}{ext}"
        import base64 as _b64
        path.write_bytes(_b64.b64decode(b64))
        return image_id
    except Exception as e:
        print(f"[image-save] 保存图片失败：{e}")
        return None


def _build_user_content_with_image_refs(
    text: str, tid: str, images: list[str]
) -> tuple[str, list[str]]:
    """处理用户消息：图片存盘 + 文本里塞 image_id 占位让 DeepSeek 知道有图。

    返回 (content_text, saved_image_ids)。content 永远是 str（不再用多模态 list），
    DeepSeek 看到占位后会主动调 vision_describe 工具向 MiMo 询问。
    """
    saved_ids: list[str] = []
    for d in images or []:
        iid = _save_image_for_conv(tid, d)
        if iid:
            saved_ids.append(iid)

    if not saved_ids:
        return text, []

    refs = "、".join(saved_ids)
    suffix = (
        f"\n\n[已上传图片：{refs}（你看不到，需要调 vision_describe(image_id, question) "
        f"让 MiMo 看图回答）]"
    )
    full = (text or "").rstrip() + suffix if text else suffix.lstrip()
    return full, saved_ids


@app.post("/api/chat")
async def chat(req: ChatRequest, request: Request):
    if req.thread_id not in conversations:
        raise HTTPException(404, "Conversation not found")
    conv = conversations[req.thread_id]

    # ── 断点续传分支：不 append user message，直接用现有 history 让 LLM 接着跑 ──
    if req.resume:
        if not conv.get("messages"):
            raise HTTPException(400, "对话还没有任何消息，无法 resume")
        # 清掉 conv 上的撞墙标记（让前端"继续"按钮消失，下次再撞会重新打上）
        conv.pop("truncated", None)
        conv.pop("truncated_reason", None)
        _checkpoint.save(conv["id"], conv)
    else:
        if conv["name"] == "新对话":
            conv["name"] = req.message[:18] + ("…" if len(req.message) > 18 else "")

        # 图片处理：存盘 + 拼 image_id 占位（DeepSeek 看占位会主动调 vision_describe）
        # 注意：再也不把图片塞进 user content 多模态 list —— 历史里永远是纯文本，
        # DeepSeek 不会再 400 "unknown variant image_url"。
        user_text, saved_image_ids = _build_user_content_with_image_refs(
            req.message, conv["id"], req.images
        )
        from datetime import datetime as _dt
        user_msg: dict[str, Any] = {
            "role": "user",
            "content": user_text,
            "ts": _dt.now().isoformat(timespec="seconds"),
        }
        if saved_image_ids:
            # UI 渲染缩略图用，不进 LLM（Message.from_dict 不读这个字段）
            user_msg["images"] = saved_image_ids
        conv["messages"].append(user_msg)
        _checkpoint.save(conv["id"], conv)
    files_before = _list_workdir_files(conv["workdir"])

    # 把 conv["messages"] 反序列化为 Message 历史喂给 Agent
    # （含上一轮的 tool_calls / reasoning_content 等完整上下文）
    history: list[Message] = [message_from_dict(m) for m in conv["messages"]]

    async def event_gen() -> AsyncIterator[dict]:
        # 模型路由：永远按对话原 model（DeepSeek 主导）。
        # 视觉能力通过 vision_describe 工具调 MiMo 实现，无需切模型。
        # master 对话会在 prompt 里注入已批准的子摘要 —— 由 get_agent_for_conv 处理。
        agent_obj = get_agent_for_conv(conv)

        # 事件合流队列：astream 的事件 + 工具内部推送的事件（如 ask_user）
        # 都走这条 queue，统一从这里 yield 给前端 SSE。
        # 这套设计让 ask_user 工具能在执行中阻塞等待用户答案，但不阻塞
        # event_gen 主循环 yield 其他事件。
        event_queue: asyncio.Queue = asyncio.Queue()

        async def event_emitter(event_dict: dict) -> None:
            """工具内部（如 ask_user）通过 config["_event_emitter"] 推事件。"""
            await event_queue.put(("custom", event_dict))

        cfg = {
            "thread_id": conv["id"],
            "workdir": conv["workdir"],
            "_event_emitter": event_emitter,
            # 子对话权限路由（tools/memory_tools.py / skills.py / self_edit.py 读取）
            "conv_kind": conv.get("kind", "standalone"),
            "sub_level": conv.get("sub_level"),
            "parent_id": conv.get("parent_id"),
        }
        new_messages: list[Message] = []
        truncated = False              # 是否撞 max_iterations 上限
        truncated_reason: str = ""

        # astream 跑在后台任务里，事件丢进 queue
        async def _astream_runner():
            try:
                async for ev in agent_obj.astream(history, config=cfg):
                    await event_queue.put(("agent", ev))
            except Exception as e:
                await event_queue.put(("exception", e))
            finally:
                await event_queue.put(("sentinel", None))

        runner_task = asyncio.create_task(_astream_runner())

        try:
            while True:
                # 客户端断开连接（点了停止按钮）→ 立即中止
                if await request.is_disconnected():
                    break
                kind, payload = await event_queue.get()
                if kind == "sentinel":
                    break
                if kind == "exception":
                    yield {
                        "event": "error",
                        "data": json.dumps({"error": f"{type(payload).__name__}: {payload}"}),
                    }
                    continue
                if kind == "custom":
                    # ask_user / command_output / command_done 等工具主动推的事件
                    yield {
                        "event": payload["type"],
                        "data": json.dumps(payload, ensure_ascii=False),
                    }
                    continue
                # kind == "agent"
                ev = payload
                etype = ev["type"]
                if etype == "delta":
                    yield {
                        "event": "delta",
                        "data": json.dumps({"text": ev["text"]}),
                    }
                elif etype == "tool_call":
                    yield {
                        "event": "tool_call",
                        "data": json.dumps({"tool": ev["name"]}),
                    }
                elif etype == "tool_result":
                    # todo_write 调用后立即推送最新清单给前端浮卡（D1）
                    if ev.get("name") == "todo_write":
                        from tools.todo import load_todos
                        yield {
                            "event": "todo_update",
                            "data": json.dumps({"items": load_todos(conv["id"])}),
                        }
                elif etype == "done":
                    new_messages = ev["new_messages"]
                    # 撞墙时 loop 也会 yield done，但带 truncated=True 标记
                    if ev.get("truncated"):
                        truncated = True
                        truncated_reason = ev.get("reason", "")
                elif etype == "error":
                    yield {
                        "event": "error",
                        "data": json.dumps({"error": ev["error"]}),
                    }
                # reasoning 暂不上抛前端（保留接口待后续展示用）
        finally:
            # 确保后台 task 被清理（正常情况它已经走完 sentinel；异常断流时取消）
            if not runner_task.done():
                runner_task.cancel()

        # 扫描新生成的文件
        files_after = _list_workdir_files(conv["workdir"])
        workdir_abs = Path(conv["workdir"]).resolve()
        new_files_rel = sorted(
            str(p.resolve().relative_to(workdir_abs)).replace("\\", "/")
            for p in (files_after - files_before)
        )

        # 把 agent 新生成的全部消息（assistant + tool_results）持久化到 conv["messages"]，
        # 同时把生成的文件附加在最后一条 assistant 消息上（前端按 role+files 渲染）
        last_assistant_idx_in_new: int | None = None
        for i, m in enumerate(new_messages):
            if m.role == "assistant":
                last_assistant_idx_in_new = i
        from datetime import datetime as _dt2
        _now_ts = _dt2.now().isoformat(timespec="seconds")
        for i, m in enumerate(new_messages):
            files = new_files_rel if i == last_assistant_idx_in_new and new_files_rel else None
            md = message_to_dict(m, files=files)
            # 新消息加时间戳（话题切分 / 时间抽屉分组用）
            md["ts"] = _now_ts
            conv["messages"].append(md)

        # 取最后一条 assistant 的文本作为前端 done 事件的 "full"
        full_text = ""
        for m in reversed(new_messages):
            if m.role == "assistant":
                full_text = m.content
                break

        # 撞墙：把标记持久化到 conv，前端切回该对话也能看到"继续"按钮
        if truncated:
            conv["truncated"] = True
            conv["truncated_reason"] = truncated_reason
        _checkpoint.save(conv["id"], conv)

        # 上下文防爆炸（C1 推广到 sub）：完成 chat 后非阻塞触发压缩检查
        # master 门槛 200，sub 门槛 150（C+ 之后调高）
        if conv.get("kind") in ("master", "sub"):
            asyncio.create_task(compress_conv_if_needed(conv))

        yield {
            "event": "done",
            "data": json.dumps({
                "full": full_text,
                "new_files": new_files_rel,
                "conv_name": conv["name"],
                "truncated": truncated,
                "truncated_reason": truncated_reason,
            }),
        }

    return EventSourceResponse(event_gen())


# ── 命令取消 ──────────────────────────────────────────────────────────────────

class CancelCommandRequest(BaseModel):
    task_id: str


@app.post("/api/command/cancel")
async def cancel_command_route(req: CancelCommandRequest):
    """取消正在运行的命令。"""
    from tools.shell import cancel_command
    ok = cancel_command(req.task_id)
    return {"ok": ok, "task_id": req.task_id}


# ── 后台任务管理器（D3）────────────────────────────────────────────────────────

from sse_starlette.sse import EventSourceResponse as _EventSourceResponse


@app.post("/api/tasks")
async def start_task(req: StartTaskRequest):
    """启动后台任务。返回 task 精简信息。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    # workdir = 请求的 workdir 或当前对话的 workdir
    workdir = req.workdir or str(DEFAULT_WORKDIR)
    task = await mgr.start(req.cmd, req.args, workdir, task_id=req.task_id)
    return task.to_dict()


@app.get("/api/tasks")
async def list_tasks():
    """列出所有后台任务。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    return [t.to_dict() for t in mgr.list()]


@app.get("/api/tasks/{task_id}")
async def get_task(task_id: str):
    """查单个任务详情（含完整输出）。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    task = mgr.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    return task.to_dict_full()


@app.post("/api/tasks/{task_id}/cancel")
async def cancel_task(task_id: str):
    """取消正在运行的任务。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    ok = mgr.cancel(task_id)
    return {"ok": ok, "task_id": task_id}


@app.get("/api/tasks/{task_id}/stream")
async def stream_task(task_id: str, request: Request):
    """SSE 实时推送任务输出行（从已缓存的行开始，新行继续推）。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    task = mgr.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")

    # 先发已缓存的行
    async def _event_gen():
        # 发当前已缓存的行
        for line in task.stdout_lines:
            yield {"event": "stdout", "data": line}
        for line in task.stderr_lines:
            yield {"event": "stderr", "data": line}

        # 如果任务已结束，不再等新行
        if task.status != "running":
            yield {"event": "done", "data": task.to_dict_full()}
            return

        # 否则轮询等新行（简单实现：每秒检查 stdout_lines 长度变化）
        sent_stdout = len(task.stdout_lines)
        sent_stderr = len(task.stderr_lines)
        import asyncio as _asyncio
        while True:
            if await request.is_disconnected():
                break
            if len(task.stdout_lines) > sent_stdout:
                for line in task.stdout_lines[sent_stdout:]:
                    yield {"event": "stdout", "data": line}
                sent_stdout = len(task.stdout_lines)
            if len(task.stderr_lines) > sent_stderr:
                for line in task.stderr_lines[sent_stderr:]:
                    yield {"event": "stderr", "data": line}
                sent_stderr = len(task.stderr_lines)
            if task.status != "running":
                yield {"event": "done", "data": task.to_dict_full()}
                break
            await _asyncio.sleep(0.5)

    return _EventSourceResponse(_event_gen())


@app.delete("/api/tasks/{task_id}")
async def delete_task(task_id: str):
    """从 registry 移除任务（不删日志）。"""
    from tools.task_manager import get_manager
    mgr = get_manager()
    ok = mgr.remove(task_id)
    return {"ok": ok}


class StartTaskRequest(BaseModel):
    cmd: str
    args: list[str] = []
    workdir: str | None = None
    task_id: str | None = None


# ── 文件上传 / 列表 / 预览 / 下载 ────────────────────────────────────────────

@app.post("/api/conversations/{tid}/upload")
async def upload_files(tid: str, files: list[UploadFile] = File(...)):
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    workdir = Path(conversations[tid]["workdir"])
    workdir.mkdir(parents=True, exist_ok=True)
    saved = []
    for f in files:
        # 安全：只取 basename
        name = Path(f.filename or "upload.bin").name
        dest = workdir / name
        # 同名冲突 → 加后缀
        if dest.exists():
            stem, suf = dest.stem, dest.suffix
            i = 1
            while dest.exists():
                dest = workdir / f"{stem}_{i}{suf}"
                i += 1
        with open(dest, "wb") as out:
            out.write(await f.read())
        saved.append(dest.name)
    return {"saved": saved}


@app.get("/api/conversations/{tid}/files")
async def list_files(tid: str, limit: int = 50):
    """列出工作目录文件 —— 已过滤 venv/__pycache__/.git/缓存目录等。

    返回最近修改的 ``limit`` 个文件（默认 50）。前端浮卡 chip 用。

    要完整树形 → ``GET /api/conversations/{tid}/files/tree``
    """
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    workdir = Path(conversations[tid]["workdir"]).resolve()
    out = []
    for rel, abs_p, is_dir in _walk_workdir(workdir):
        if is_dir:
            continue
        try:
            stat = abs_p.stat()
        except OSError:
            continue
        out.append({"path": rel, "size": stat.st_size, "mtime": stat.st_mtime})
    out.sort(key=lambda f: f["mtime"], reverse=True)
    return out[:limit] if limit > 0 else out


@app.get("/api/conversations/{tid}/files/tree")
async def list_files_tree(tid: str):
    """工作目录树形结构（已过滤 venv 等）。

    每个节点形如:
      {"name": "...", "path": "rel/path", "is_dir": bool,
       "size": int, "mtime": float, "children": [...]}

    目录按字典序，文件在目录之后；空目录不返回。
    """
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    workdir = Path(conversations[tid]["workdir"]).resolve()
    if not workdir.exists():
        return []

    # 收集 (rel, abs_path, is_dir, stat) 列表
    nodes_by_path: dict[str, dict] = {}
    for rel, abs_p, is_dir in _walk_workdir(workdir):
        try:
            stat = abs_p.stat()
        except OSError:
            continue
        nodes_by_path[rel] = {
            "name": abs_p.name,
            "path": rel,
            "is_dir": is_dir,
            "size": 0 if is_dir else stat.st_size,
            "mtime": stat.st_mtime,
            "children": [] if is_dir else None,
        }

    # 把每个节点挂到父节点的 children 上
    roots: list[dict] = []
    for rel, node in nodes_by_path.items():
        parent_rel = "/".join(rel.split("/")[:-1])
        if parent_rel and parent_rel in nodes_by_path:
            nodes_by_path[parent_rel]["children"].append(node)
        else:
            roots.append(node)

    def _sort_recursive(nodes: list[dict]):
        nodes.sort(key=lambda n: (not n["is_dir"], n["name"].lower()))  # 目录在前
        for n in nodes:
            if n["is_dir"] and n["children"]:
                _sort_recursive(n["children"])

    _sort_recursive(roots)
    return roots


@app.post("/api/conversations/{tid}/open")
async def open_workdir(tid: str):
    """用系统文件管理器打开工作目录（Windows explorer / Mac open / Linux xdg-open）。"""
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    workdir = Path(conversations[tid]["workdir"]).resolve()
    if not workdir.exists():
        try:
            workdir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            raise HTTPException(500, f"工作目录不存在且无法创建: {e}")

    import subprocess
    try:
        if sys.platform == "win32":
            # explorer.exe 返回非 0 也算成功，所以直接 Popen 不 check
            subprocess.Popen(["explorer.exe", str(workdir)])
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(workdir)])
        else:
            subprocess.Popen(["xdg-open", str(workdir)])
        return {"ok": True, "path": str(workdir)}
    except Exception as e:
        raise HTTPException(500, f"打开工作目录失败: {e}")


@app.get("/api/conversations/{tid}/preview")
async def preview_file(tid: str, path: str):
    """提取 docx / pdf 文件的纯文本预览（前端展开预览用）。

    返回 {"text": str, "truncated": bool}；不支持的类型返回 {"error": ...}。
    """
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    target = _safe_path(conversations[tid]["workdir"], path)
    if not target.exists() or not target.is_file():
        raise HTTPException(404, "File not found")

    MAX_CHARS = 12000  # 前端 modal 显示，比 read_file 工具略多
    ext = target.suffix.lower()
    try:
        if ext == ".pdf":
            import pdfplumber
            parts = []
            with pdfplumber.open(target) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text() or "(无文字)"
                    parts.append(f"=== Page {i} ===\n{t}")
            text = "\n\n".join(parts)
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(target))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts) or "(空文档)"
        else:
            return JSONResponse({"error": f"unsupported extension: {ext}"}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"{type(e).__name__}: {e}"}, status_code=500)

    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS] + f"\n\n...(已截断，原文 {len(text)} 字符)"
    return {"text": text, "truncated": truncated}


@app.get("/api/conversations/{tid}/file")
async def get_file(tid: str, path: str, download: bool = False):
    """读取工作目录内文件，用于预览（<img>）或下载。"""
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    target = _safe_path(conversations[tid]["workdir"], path)
    if not target.exists() or not target.is_file():
        raise HTTPException(404, "File not found")
    mime, _ = mimetypes.guess_type(target.name)
    if download:
        return FileResponse(target, filename=target.name, media_type=mime or "application/octet-stream")
    return FileResponse(target, media_type=mime or "application/octet-stream")


@app.get("/api/conversations/{tid}/image/{image_id}")
async def get_chat_image(tid: str, image_id: str):
    """返回某对话历史里上传的图片（前端消息渲染时 <img src> 用）。

    存放位置：``.sandbox/_meta/<tid>/images/<image_id>.<ext>``
    image_id 必须是合法标识符（防路径穿越），扩展名由实际文件决定。
    """
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    # 防御性：image_id 必须是 img_xxxxxxxx 格式
    if not image_id.replace("_", "").isalnum() or len(image_id) > 32:
        raise HTTPException(400, "Invalid image_id")
    from tools.vision import find_image_path
    path = find_image_path(tid, image_id)
    if path is None or not path.exists():
        raise HTTPException(404, "Image not found")
    mime, _ = mimetypes.guess_type(path.name)
    return FileResponse(path, media_type=mime or "application/octet-stream")


# ── 长期记忆 API ──────────────────────────────────────────────────────────────


class AddMemoryRequest(BaseModel):
    fact: str
    category: str = "other"
    importance: int = 5


class UpdateMemoryRequest(BaseModel):
    text: str | None = None
    category: str | None = None
    importance: int | None = None


@app.get("/api/memory")
async def list_memory():
    from memory import list_memories
    return list_memories()


@app.post("/api/memory")
async def add_memory_route(req: AddMemoryRequest):
    from memory import add_memory
    try:
        mem_id = add_memory(req.fact, category=req.category, importance=req.importance)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {
        "id": mem_id,
        "fact": req.fact,
        "category": req.category,
        "importance": req.importance,
    }


class MemorySettingsRequest(BaseModel):
    memory_write_enabled: bool | None = None


# 注意：/api/memory/settings 必须声明在 /api/memory/{mem_id} 之前，
# 否则 FastAPI 会按声明顺序先匹配 {mem_id}，把 "settings" 当成 id 处理。
@app.get("/api/memory/settings")
async def get_memory_settings():
    from memory import load_settings
    return load_settings()


@app.patch("/api/memory/settings")
async def update_memory_settings(req: MemorySettingsRequest):
    from memory import save_settings
    payload = {}
    if req.memory_write_enabled is not None:
        payload["memory_write_enabled"] = req.memory_write_enabled
    return save_settings(payload)


@app.get("/api/memory/{mem_id}")
async def get_memory_by_id(mem_id: str):
    """按 id 取单条记忆全文 —— 前端"已压缩话题"折叠卡展开时用。

    ChromaDB Python SDK 没有内建 by_id 查询，用 collection.get(ids=[...]) 拿。
    """
    from memory import _get_collection
    coll = _get_collection()
    try:
        result = coll.get(ids=[mem_id])
    except Exception as e:
        raise HTTPException(500, f"读取记忆失败: {e}")
    if not result or not result.get("ids"):
        raise HTTPException(404, "Memory not found")
    docs = result.get("documents") or []
    metas = result.get("metadatas") or []
    if not docs:
        raise HTTPException(404, "Memory not found")
    meta = metas[0] if metas else {}
    return {
        "id": mem_id,
        "text": docs[0],
        "category": (meta or {}).get("category", "other"),
        "importance": int((meta or {}).get("importance", 5)),
        "created_at": (meta or {}).get("created_at", ""),
    }


@app.patch("/api/memory/{mem_id}")
async def update_memory_route(mem_id: str, req: UpdateMemoryRequest):
    from memory import update_memory
    try:
        return update_memory(
            mem_id,
            text=req.text,
            category=req.category,
            importance=req.importance,
        )
    except KeyError as e:
        raise HTTPException(404, str(e))
    except ValueError as e:
        raise HTTPException(400, str(e))


@app.delete("/api/memory/{mem_id}")
async def delete_memory_route(mem_id: str):
    from memory import delete_memory
    delete_memory(mem_id)
    return {"ok": True}


# ── 主对话防爆炸（M-4）─────────────────────────────────────────────────────
#
# 主对话 messages 长了会撑爆 token + 历史变重。M-4 设计：
# 1) 按"时间间隔 > 2h"切分 messages 为话题段（topic）
# 2) 触发条件：messages > 100 条 OR 距上次整理 > 24h
# 3) 把最旧的 N 个话题段（除最近 1 个）单独跑 LLM 摘要 → 写进
#    memory（category=chat_log）
# 4) master.messages 砍掉这些已压缩的话题段
#
# chat_log 类别在 recall 时默认被过滤（避免污染严肃检索），但是有希查
# "我们上次聊到 X" 时显式包含。


COMPRESS_GAP_MINUTES = 120              # 话题切分阈值：≥ 2h 间隔视为新话题
COMPRESS_MSG_THRESHOLD_MASTER = 200     # master 触发门槛（提到 200 配合 KEEP_TOPICS=5，给 5 个 topic 留空间）
COMPRESS_MSG_THRESHOLD_SUB = 150        # sub 触发门槛（同步上调）
COMPRESS_HOURS_THRESHOLD = 24           # 触发整理的距上次时间下限（小时）
COMPRESS_KEEP_TOPICS = 5                # 至少保留最近 N 个话题（视觉上还能看到这 N 个 topic 的完整消息）
COMPRESS_BACKUP_KEEP_DAYS = 7           # .pre-compress-<ts>.bak 保留天数（自动清理）

# legacy 别名（旧代码/测试仍可用）
COMPRESS_MSG_THRESHOLD = COMPRESS_MSG_THRESHOLD_MASTER


def _split_into_topics(messages: list[dict]) -> list[list[int]]:
    """按时间戳间隔切分 messages，返回话题段的索引列表。

    例：返回 ``[[0,1,2], [3,4,5,6], [7,8]]`` 表示三个话题段。
    没 ts 的消息默认接到上一段（保守，不强制切分）。
    """
    from datetime import datetime as _dt
    segments: list[list[int]] = []
    current: list[int] = []
    last_ts: float | None = None
    for i, m in enumerate(messages):
        ts_str = m.get("ts")
        if ts_str:
            try:
                ts = _dt.fromisoformat(ts_str).timestamp()
                if last_ts is not None and ts - last_ts > COMPRESS_GAP_MINUTES * 60:
                    if current:
                        segments.append(current)
                    current = []
                last_ts = ts
            except Exception:
                pass
        current.append(i)
    if current:
        segments.append(current)
    return segments


def _should_compress_conv(conv: dict) -> bool:
    """检查 master / sub 对话是否到该整理的时刻。

    master 门槛 100；sub 门槛 80（任务型对话容易高频堆积，更早压缩降低 LLM 退化风险）。
    standalone 不压缩（兼容老行为）。
    """
    kind = conv.get("kind")
    if kind not in ("master", "sub"):
        return False
    threshold = (
        COMPRESS_MSG_THRESHOLD_MASTER if kind == "master"
        else COMPRESS_MSG_THRESHOLD_SUB
    )
    msgs = conv.get("messages") or []
    if len(msgs) > threshold:
        return True
    last = conv.get("last_compress_at", "")
    if not last:
        return len(msgs) > 30  # 没整理过且消息够多
    from datetime import datetime as _dt, timedelta as _td
    try:
        last_dt = _dt.fromisoformat(last)
        return _dt.now() - last_dt > _td(hours=COMPRESS_HOURS_THRESHOLD) and len(msgs) > 50
    except Exception:
        return False


def _should_compress_master(conv: dict) -> bool:
    """legacy 别名 —— 仅 master 触发。新代码用 _should_compress_conv。"""
    return conv.get("kind") == "master" and _should_compress_conv(conv)


async def _summarize_topic(topic_msgs: list[dict]) -> str:
    """单独跑 LLM 给一段话题做摘要。复用 _generate_conv_summary 的策略。"""
    fake_conv = {"messages": topic_msgs}
    return await _generate_conv_summary(fake_conv)


def _backup_conv_before_compress(conv_id: str) -> Path | None:
    """压缩前把当前 conv.json 复制成 .pre-compress-<ts>.bak。

    7 天后自动清理。备份永不阻塞主流程（任何失败静默返回 None）。
    """
    try:
        import shutil
        from datetime import datetime as _dt, timedelta as _td
        src = _checkpoint._path(conv_id)
        if not src.exists():
            return None
        ts = _dt.now().strftime("%Y%m%d_%H%M%S")
        dst = src.parent / f"conv.json.pre-compress-{ts}.bak"
        shutil.copy2(src, dst)

        # 清理 7 天前的 .bak（每次压缩做一次，省个 cron）
        cutoff = _dt.now() - _td(days=COMPRESS_BACKUP_KEEP_DAYS)
        for old in src.parent.glob("conv.json.pre-compress-*.bak"):
            try:
                if _dt.fromtimestamp(old.stat().st_mtime) < cutoff:
                    old.unlink()
            except Exception:
                pass
        return dst
    except Exception as e:
        print(f"[compress] backup 失败（不阻塞）: {e}")
        return None


async def compress_conv_if_needed(conv: dict) -> dict:
    """如果需要，把 master/sub 的老话题段压缩进 chat_log 记忆。

    chat_log 文本里包含来源标识（``[主对话/sub:<name>] 片段...``），
    recall 时能定位是哪个对话的旧片段。

    每次压缩：
    1. **先备份** conv.json → ``conv.json.pre-compress-<ts>.bak``（7 天保留）
    2. 把超出 ``COMPRESS_KEEP_TOPICS`` 的旧 topic 一段段送 LLM 摘要 → 写记忆库
    3. ``conv["compressed_summaries"]`` append 每段 entry（前端可显示折叠卡）

    返回 ``{"compressed": int, "kept_messages": int}``。失败 / 不需要时
    ``compressed=0``。所有失败永不阻塞主流程。
    """
    summary = {"compressed": 0, "kept_messages": len(conv.get("messages") or [])}
    try:
        if not _should_compress_conv(conv):
            return summary

        msgs = conv.get("messages") or []
        topics = _split_into_topics(msgs)
        if len(topics) <= COMPRESS_KEEP_TOPICS:
            return summary

        # ── 1. 备份原 conv.json（数据安全）──
        _backup_conv_before_compress(conv["id"])

        to_compress = topics[:-COMPRESS_KEEP_TOPICS]
        to_keep = [idx for tp in topics[-COMPRESS_KEEP_TOPICS:] for idx in tp]

        # 来源标识（chat_log 文本里包含，方便 recall 定位）
        conv_kind = conv.get("kind", "?")
        conv_name = conv.get("name", "?")
        if conv_kind == "master":
            source_label = "主对话"
        elif conv_kind == "sub":
            source_label = f"子对话「{conv_name}」"
        else:
            source_label = conv_name

        from datetime import datetime as _dt
        compress_ts = _dt.now().isoformat(timespec="seconds")
        new_summary_entries: list[dict] = []
        compressed_count = 0
        for topic_indices in to_compress:
            topic_msgs = [msgs[i] for i in topic_indices]
            if len(topic_msgs) < 2:
                continue
            topic_summary = await _summarize_topic(topic_msgs)
            if not topic_summary or topic_summary.startswith("[摘要失败"):
                to_keep.extend(topic_indices)
                continue

            first_ts = topic_msgs[0].get("ts", "") if topic_msgs else ""
            last_ts = topic_msgs[-1].get("ts", "") if topic_msgs else ""
            text_to_save = (
                f"[{source_label} 片段 {first_ts[:10]} 至 {last_ts[:10]}，"
                f"共 {len(topic_msgs)} 条消息]\n{topic_summary}"
            )

            try:
                from memory import add_memory
                mem_id = add_memory(text_to_save, category="chat_log", importance=3)
                compressed_count += 1
                # ── 3. 给前端折叠卡用的 entry ──
                new_summary_entries.append({
                    "ts": compress_ts,
                    "first_ts": first_ts,
                    "last_ts": last_ts,
                    "count": len(topic_msgs),
                    "memory_id": mem_id,
                    "summary_preview": topic_summary[:240],
                })
            except Exception as e:
                print(f"[compress] 写 chat_log 失败: {e}")
                to_keep.extend(topic_indices)

        to_keep.sort()
        new_messages = [msgs[i] for i in to_keep]
        conv["messages"] = new_messages
        conv["last_compress_at"] = compress_ts
        # append 而不是覆盖：多次压缩历史累积
        existing = list(conv.get("compressed_summaries") or [])
        existing.extend(new_summary_entries)
        conv["compressed_summaries"] = existing
        _checkpoint.save(conv["id"], conv)

        summary["compressed"] = compressed_count
        summary["kept_messages"] = len(new_messages)
        if compressed_count:
            print(
                f"[compress] {source_label} 压缩 {compressed_count} 个话题段进 chat_log，"
                f"messages: {len(msgs)} → {len(new_messages)}"
            )
    except Exception as e:
        print(f"[compress] 异常（不影响主流程）: {e}")
    return summary


async def compress_master_if_needed(conv: dict) -> dict:
    """legacy 别名 —— 仅 master 触发。新代码用 compress_conv_if_needed。"""
    if conv.get("kind") != "master":
        return {"compressed": 0, "kept_messages": len(conv.get("messages") or [])}
    return await compress_conv_if_needed(conv)


@app.post("/api/conversations/{tid}/compress")
async def compress_conv_endpoint(tid: str):
    """手动触发对话压缩（master / sub 都可，前端"立即整理"按钮用）。"""
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    conv = conversations[tid]
    if conv.get("kind") not in ("master", "sub"):
        raise HTTPException(400, "只有 master / sub 对话可以压缩，standalone 不行")
    result = await compress_conv_if_needed(conv)
    return result


# ── 子对话摘要 API（M-3）────────────────────────────────────────────────────


async def _generate_conv_summary(conv: dict) -> str:
    """调 LLM 生成子对话的话题摘要。

    取最近的 user/assistant 消息（剔除 tool 结果 / 图片占位），单独跑一次
    DeepSeek 生成 200-400 字摘要。失败返回空字符串。
    """
    msgs = conv.get("messages") or []
    text_lines: list[str] = []
    for m in msgs:
        role = m.get("role", "")
        if role not in ("user", "assistant"):
            continue
        content = m.get("content", "")
        # 多模态 content 取 text part
        if isinstance(content, list):
            text_parts = [p.get("text", "") for p in content if isinstance(p, dict) and p.get("type") == "text"]
            content = " ".join(text_parts)
        if not isinstance(content, str):
            continue
        content = content.strip()
        if not content:
            continue
        # 去掉图片占位（vision 路由生成的）
        if "[已上传图片：" in content:
            import re as _re
            content = _re.sub(r"\[已上传图片：[^\]]*\]", "", content).strip()
        if content:
            text_lines.append(f"{role}: {content[:600]}")

    if len(text_lines) < 2:
        return ""

    # 取最近 80 段（防 prompt 太长）
    conversation_text = "\n\n".join(text_lines[-80:])

    summary_prompt = (
        "请用一段简短的中文（200-400 字）总结下面这段对话的核心内容。\n"
        "格式要求：\n"
        "1. 第一段（1-2 句）：主要话题是什么\n"
        "2. 第二段（2-4 句）：关键结论 / 产出 / 数据\n"
        "3. 第三段（可选，1-2 句）：未完成的待办或主人下次会继续的事\n\n"
        "不要总结纯闲聊 / 寒暄 / 礼貌用语，只写实质内容。"
        "不要复述对话过程，只要结果。\n\n"
        f"对话内容：\n{conversation_text}"
    )

    try:
        from ai_agent import DeepSeekClient, Message
    except Exception as e:
        return f"[摘要失败：import LLM 客户端 - {e}]"

    try:
        client = DeepSeekClient(model="deepseek-v4-flash", temperature=0.3)
    except ValueError as e:
        return f"[摘要失败：{e}]"

    parts: list[str] = []
    err_msg: str | None = None
    try:
        async for ev in client.stream([Message.user(summary_prompt)]):
            if ev.get("type") == "delta":
                parts.append(ev.get("text", ""))
            elif ev.get("type") == "error":
                err_msg = ev.get("error", "未知")
                break
    except Exception as e:
        err_msg = f"{type(e).__name__}: {e}"
    finally:
        try:
            await client.aclose()
        except Exception:
            pass

    if err_msg:
        return f"[摘要失败：{err_msg}]"

    return "".join(parts).strip()


@app.post("/api/conversations/{tid}/summarize")
async def summarize_conv(tid: str):
    """生成 / 重新生成子对话的摘要。生成后 ``summary_pending_approval=True``，
    主人需要 PATCH ``summary_approved_for_master`` 才能纳入主对话 prompt。
    """
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    conv = conversations[tid]
    if conv.get("kind") != "sub":
        raise HTTPException(400, "只有子对话可生成摘要")

    summary = await _generate_conv_summary(conv)
    if not summary or summary.startswith("[摘要失败"):
        return {"ok": False, "reason": summary or "对话内容过少，无法摘要"}

    from datetime import datetime as _dt
    conv["summary"] = summary
    conv["summary_updated_at"] = _dt.now().isoformat(timespec="seconds")
    conv["summary_pending_approval"] = True
    conv["summary_approved_for_master"] = False
    _checkpoint.save(tid, conv)

    return {
        "ok": True,
        "summary": summary,
        "summary_updated_at": conv["summary_updated_at"],
    }


class SummaryApprovalRequest(BaseModel):
    approved: bool


@app.patch("/api/conversations/{tid}/summary")
async def approve_summary(tid: str, req: SummaryApprovalRequest):
    """主人批准（或撤销批准）子对话的摘要纳入主对话 prompt。"""
    if tid not in conversations:
        raise HTTPException(404, "Conversation not found")
    conv = conversations[tid]
    if conv.get("kind") != "sub":
        raise HTTPException(400, "只有子对话有摘要可审批")
    if not (conv.get("summary") or "").strip():
        raise HTTPException(400, "该子对话还没有摘要，先调 POST /summarize")
    conv["summary_approved_for_master"] = req.approved
    conv["summary_pending_approval"] = False
    _checkpoint.save(tid, conv)
    return {
        "ok": True,
        "summary_approved_for_master": conv["summary_approved_for_master"],
    }


# ── Todo 列表 API（D1）─────────────────────────────────────────────────────────


@app.get("/api/conversations/{tid}/todos")
async def get_todos(tid: str):
    """读取某个对话的 todo 清单（前端切换对话时拉一次）。"""
    from tools.todo import load_todos
    return {"items": load_todos(tid)}


# ── ask_user 答案回传 API（D2）─────────────────────────────────────────────────


class AskUserAnswer(BaseModel):
    answer: str


@app.post("/api/ask_user/{ask_id}")
async def submit_ask_user_answer(ask_id: str, payload: AskUserAnswer):
    """前端把用户的回答回传给等待中的 ask_user 工具。

    工具内部 await 的 Future 会被 set_result，工具继续执行返回给 agent。
    如果 ask_id 不存在（超时或刷新页面），返回 410 让前端知道弹窗失效。
    """
    from tools.dialog import get_pending_future
    fut = get_pending_future(ask_id)
    if fut is None or fut.done():
        raise HTTPException(410, "该提问已超时或已被回答")
    fut.set_result(payload.answer)
    return {"ok": True}


# ── 健康检查 ─────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    from memory import count_memories
    try:
        mem_count = count_memories()
    except Exception:
        mem_count = -1
    return {
        "status": "ok",
        "models": list(MODELS),
        "conversations": len(conversations),
        "memories": mem_count,
    }


# ── 入口 ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import socket
    import threading
    import time
    import webbrowser
    import uvicorn

    PORT = 3616
    URL = f"http://127.0.0.1:{PORT}"

    def _open_browser_when_ready():
        """轮询端口直到 uvicorn 真正接受连接，再唤醒浏览器；最多等 10 秒。"""
        deadline = time.time() + 10
        while time.time() < deadline:
            try:
                with socket.create_connection(("127.0.0.1", PORT), timeout=0.3):
                    webbrowser.open(URL)
                    return
            except OSError:
                time.sleep(0.2)

    threading.Thread(target=_open_browser_when_ready, daemon=True).start()
    uvicorn.run("server:app", host="127.0.0.1", port=PORT, reload=False)
